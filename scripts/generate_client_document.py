#!/usr/bin/env python3
"""Generate a client-facing health interpretation and 1v1 plan document.

The input is an internal JSON payload assembled by the agent. The client does
not need to prepare or copy this payload. The generator produces a styled DOCX
and an HTML preview in the requested output directory.
"""

from __future__ import annotations

import argparse
import contextlib
import html
import io
import json
import re
import sys
from datetime import date
from pathlib import Path
from typing import Any


try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - helpful CLI failure
    raise SystemExit(
        "缺少 python-docx，请先运行：.venv-docs/bin/python -m pip install -r "
        "scripts/requirements-docs.txt"
    ) from exc


INK = "203238"
MUTED = "5F7074"
TEAL = "1D6B6A"
TEAL_DARK = "154B4C"
MINT = "E9F3F0"
CORAL = "C65B48"
SAND = "F7F5F0"
LINE = "D7E1DF"
WHITE = "FFFFFF"
FONT = "PingFang SC"


def text(value: Any) -> str:
    return str(value).strip() if value is not None else ""


def list_value(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value.strip()] if value.strip() else []
    return [text(item) for item in value if text(item)]


def safe_name(value: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|\n\r\t]+", "", value).strip()
    return name or "客户"


def esc(value: Any) -> str:
    return html.escape(text(value))


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 120, start: int = 150, bottom: int = 120, end: int = 150) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: float | None = None, color: str = INK, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def style_document(document: Document, brand: str, report_date: str) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for style_name, size, color in (("Title", 29, TEAL_DARK), ("Heading 1", 19, TEAL_DARK), ("Heading 2", 13, TEAL)):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{brand}  ·  {report_date}")
    set_run_font(run, 8.5, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("客户专属健康管理资料  ·  ")
    set_run_font(run, 8, MUTED)
    add_field(footer, "PAGE")


def add_text_paragraph(document: Document, content: str, *, style: str | None = None, color: str = INK, size: float | None = None, bold: bool = False, align=None):
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text(content))
    set_run_font(run, size, color, bold)
    return paragraph


def add_bullets(document: Document, items: list[str], *, color: str = INK) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Mm(5)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_run_font(run, 10.5, color)


def add_callout(document: Document, heading: str, content: str, fill: str = MINT, accent: str = TEAL) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_border(cell, fill, "0")
    set_cell_margins(cell, 170, 210, 170, 210)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(heading)
    set_run_font(run, 10.5, accent, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(content)
    set_run_font(run, 10.5, INK)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_header(document: Document, number: str, title: str, lead: str = "") -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    badge = p.add_run(f"{number}  ")
    set_run_font(badge, 11, CORAL, True)
    heading = p.add_run(title)
    set_run_font(heading, 19, TEAL_DARK, True)
    if lead:
        add_text_paragraph(document, lead, color=MUTED, size=10.5)


def add_cover(document: Document, data: dict[str, Any], report_date: str) -> None:
    for _ in range(3):
        document.add_paragraph()
    brand = text(data.get("brand"))
    customer = text(data.get("customer_name"))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(brand or "健康管理服务")
    set_run_font(run, 12, CORAL, True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("体检报告解读\n与 1v1 健康管理方案")
    set_run_font(run, 28, TEAL_DARK, True)
    p.paragraph_format.space_after = Pt(13)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"为 {customer} 定制")
    set_run_font(run, 13, MUTED)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report_date)
    set_run_font(run, 10, MUTED)
    document.add_paragraph()
    add_callout(document, "这份报告想解决的事", text(data.get("cover_subtitle")) or "看懂身体现在的状态，并找到真正做得下去的下一步。", fill=SAND, accent=CORAL)
    for _ in range(5):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text(data.get("advisor")) or "健康管理顾问")
    set_run_font(run, 10, MUTED)
    document.add_page_break()


def add_summary(document: Document, summary: list[str], profile: list[dict[str, Any]] | None = None) -> None:
    add_section_header(document, "01", "先看结论", "先把最重要的判断放在前面，避免被体检单上的一堆箭头带偏。")
    add_profile(document, profile or [])
    if summary:
        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for index, item in enumerate(summary):
            row = table.add_row()
            left, right = row.cells
            set_cell_fill(left, TEAL if index == 0 else MINT)
            set_cell_fill(right, SAND)
            for cell in (left, right):
                set_cell_border(cell, WHITE, "0")
                set_cell_margins(cell, 150, 180, 150, 180)
            label = left.paragraphs[0]
            r = label.add_run(f"重点 {index + 1}")
            set_run_font(r, 10, WHITE if index == 0 else TEAL_DARK, True)
            content = right.paragraphs[0]
            r = content.add_run(item)
            set_run_font(r, 10.5, INK)
        document.add_paragraph()


def add_profile(document: Document, profile: list[dict[str, Any]]) -> None:
    if not profile:
        return
    add_text_paragraph(document, "客户摘要", style="Heading 2")
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index in range(0, len(profile), 2):
        row = table.add_row()
        pair = profile[index : index + 2]
        for column, item in enumerate(pair):
            cell = row.cells[column]
            set_cell_fill(cell, SAND if (index + column) % 2 == 0 else MINT)
            set_cell_border(cell, WHITE, "0")
            set_cell_margins(cell, 140, 170, 140, 170)
            paragraph = cell.paragraphs[0]
            label = paragraph.add_run(text(item.get("label")) + "\n")
            set_run_font(label, 8.5, MUTED, True)
            value = paragraph.add_run(text(item.get("value")))
            set_run_font(value, 10.5, INK, True)
        if len(pair) == 1:
            set_cell_fill(row.cells[1], WHITE)
            set_cell_border(row.cells[1], WHITE, "0")
    document.add_paragraph()


def add_metrics(document: Document, metrics: list[dict[str, Any]]) -> None:
    if not metrics:
        return
    add_text_paragraph(document, "本次血脂结果", style="Heading 2")
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, ("指标", "结果", "报告提示", "怎么理解")):
        set_cell_fill(cell, TEAL_DARK)
        set_cell_border(cell, TEAL_DARK, "0")
        set_cell_margins(cell, 120, 120, 120, 120)
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, 9.5, WHITE, True)
    for index, item in enumerate(metrics):
        row = table.add_row()
        values = (text(item.get("name")), text(item.get("value")), text(item.get("status")), text(item.get("meaning")))
        for cell, value in zip(row.cells, values):
            set_cell_fill(cell, SAND if index % 2 == 0 else WHITE)
            set_cell_border(cell, LINE, "4")
            set_cell_margins(cell, 120, 120, 120, 120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, 9, INK)
    document.add_paragraph()


def add_report_content(document: Document, report: dict[str, Any]) -> None:
    add_section_header(document, "02", "体检报告解读", text(report.get("intro")))
    if report.get("conclusion"):
        add_callout(document, "一句话判断", text(report["conclusion"]), fill=MINT, accent=TEAL)
    add_metrics(document, report.get("metrics", []))

    report_sections = report.get("sections", [])
    for section in report_sections:
        title = text(section.get("title"))
        if title:
            add_text_paragraph(document, title, style="Heading 2")
        if section.get("lead"):
            add_callout(document, "核心判断", text(section["lead"]), fill=SAND, accent=CORAL)
        for paragraph in list_value(section.get("paragraphs")):
            add_text_paragraph(document, paragraph)
        add_bullets(document, list_value(section.get("bullets")))

    if report.get("risks"):
        add_text_paragraph(document, "这些问题可能带来的现实影响", style="Heading 2")
        add_bullets(document, list_value(report["risks"]))
    if report.get("actions"):
        add_text_paragraph(document, "接下来最该做的事", style="Heading 2")
        add_bullets(document, list_value(report["actions"]))
    if report.get("closing"):
        add_callout(document, "把报告压成一句话", text(report["closing"]), fill=MINT, accent=TEAL)


def add_plan_content(document: Document, plan: dict[str, Any]) -> None:
    document.add_page_break()
    add_section_header(document, "03", "1v1 定制方案", text(plan.get("intro")))
    if plan.get("overview"):
        add_callout(document, "你的方案主线", text(plan["overview"]), fill=MINT, accent=TEAL)
    if plan.get("goals"):
        add_text_paragraph(document, "阶段目标", style="Heading 2")
        add_bullets(document, list_value(plan["goals"]))

    phases = plan.get("phases", [])
    if phases:
        add_text_paragraph(document, "执行节奏", style="Heading 2")
        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ("阶段", "重点", "你要做什么")
        for cell, header in zip(table.rows[0].cells, headers):
            set_cell_fill(cell, TEAL_DARK)
            set_cell_border(cell, TEAL_DARK, "0")
            set_cell_margins(cell, 130, 150, 130, 150)
            run = cell.paragraphs[0].add_run(header)
            set_run_font(run, 10, WHITE, True)
        for index, phase in enumerate(phases):
            row = table.add_row()
            values = (text(phase.get("name")) or f"第 {index + 1} 阶段", text(phase.get("focus")), "\n".join(list_value(phase.get("actions"))))
            for cell, value in zip(row.cells, values):
                set_cell_fill(cell, SAND if index % 2 == 0 else WHITE)
                set_cell_border(cell, LINE, "4")
                set_cell_margins(cell, 130, 150, 130, 150)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                run = cell.paragraphs[0].add_run(value)
                set_run_font(run, 9.5, INK)
        document.add_paragraph()

    for title, key in (("日常执行清单", "daily_checklist"), ("本阶段重点提醒", "reminders"), ("复查与随访", "follow_up")):
        values = list_value(plan.get(key))
        if values:
            add_text_paragraph(document, title, style="Heading 2")
            add_bullets(document, values)


def add_notes(document: Document, data: dict[str, Any]) -> None:
    notes = list_value(data.get("notes"))
    disclaimer = text(data.get("disclaimer")) or "以上内容用于健康管理沟通，不构成诊疗建议。如有具体健康问题，请线下就医。"
    if notes:
        add_text_paragraph(document, "说明", style="Heading 2")
        add_bullets(document, notes, color=MUTED)
    add_callout(document, "重要说明", disclaimer, fill=SAND, accent=CORAL)


def build_docx(data: dict[str, Any], output_path: Path, report_date: str) -> None:
    document = Document()
    style_document(document, text(data.get("brand")) or "健康管理服务", report_date)
    add_cover(document, data, report_date)
    add_summary(document, list_value(data.get("summary")), data.get("profile", []))
    add_report_content(document, data.get("report", {}))
    add_plan_content(document, data.get("plan", {}))
    add_notes(document, data)
    document.save(output_path)


def html_list(items: list[str]) -> str:
    if not items:
        return ""
    return "<ul>" + "".join(f"<li>{esc(item)}</li>" for item in items) + "</ul>"


def html_callout(heading: str, content: str, tone: str = "mint") -> str:
    return f'<div class="callout {tone}"><strong>{esc(heading)}</strong><p>{esc(content)}</p></div>'


def build_html(data: dict[str, Any], output_path: Path, report_date: str) -> None:
    brand = text(data.get("brand")) or "健康管理服务"
    customer = text(data.get("customer_name")) or "客户"
    summary = list_value(data.get("summary"))
    profile = data.get("profile", [])
    report = data.get("report", {})
    plan = data.get("plan", {})

    summary_html = "".join(f'<div class="summary-row"><span>重点 {i + 1}</span><p>{esc(item)}</p></div>' for i, item in enumerate(summary))
    profile_html = "".join(
        f'<div class="profile-item"><span>{esc(item.get("label"))}</span><strong>{esc(item.get("value"))}</strong></div>'
        for item in profile
    )
    profile_block = f'<h3>客户摘要</h3><div class="profile">{profile_html}</div>' if profile_html else ""
    report_html = ""
    if report.get("conclusion"):
        report_html += html_callout("一句话判断", report["conclusion"])
    metrics = report.get("metrics", [])
    if metrics:
        rows = "".join(
            f"<tr><td>{esc(item.get('name'))}</td><td>{esc(item.get('value'))}</td><td>{esc(item.get('status'))}</td><td>{esc(item.get('meaning'))}</td></tr>"
            for item in metrics
        )
        report_html += "<h3>本次血脂结果</h3><table><thead><tr><th>指标</th><th>结果</th><th>报告提示</th><th>怎么理解</th></tr></thead><tbody>" + rows + "</tbody></table>"
    for section in report.get("sections", []):
        report_html += f'<h3>{esc(section.get("title"))}</h3>'
        if section.get("lead"):
            report_html += html_callout("核心判断", section["lead"], "sand")
        report_html += "".join(f"<p>{esc(item)}</p>" for item in list_value(section.get("paragraphs")))
        report_html += html_list(list_value(section.get("bullets")))
    if report.get("risks"):
        report_html += f'<h3>这些问题可能带来的现实影响</h3>{html_list(list_value(report["risks"]))}'
    if report.get("actions"):
        report_html += f'<h3>接下来最该做的事</h3>{html_list(list_value(report["actions"]))}'
    if report.get("closing"):
        report_html += html_callout("把报告压成一句话", report["closing"])

    phases_html = ""
    for phase in plan.get("phases", []):
        actions = "".join(f"<li>{esc(item)}</li>" for item in list_value(phase.get("actions")))
        phases_html += f"<tr><td>{esc(phase.get('name'))}</td><td>{esc(phase.get('focus'))}</td><td><ul>{actions}</ul></td></tr>"
    plan_html = ""
    if plan.get("overview"):
        plan_html += html_callout("你的方案主线", plan["overview"])
    if plan.get("goals"):
        plan_html += f'<h3>阶段目标</h3>{html_list(list_value(plan["goals"]))}'
    if phases_html:
        plan_html += '<h3>执行节奏</h3><table><thead><tr><th>阶段</th><th>重点</th><th>你要做什么</th></tr></thead><tbody>' + phases_html + "</tbody></table>"
    for title, key in (("日常执行清单", "daily_checklist"), ("本阶段重点提醒", "reminders"), ("复查与随访", "follow_up")):
        values = list_value(plan.get(key))
        if values:
            plan_html += f"<h3>{title}</h3>{html_list(values)}"

    notes_html = html_list(list_value(data.get("notes")))
    disclaimer = text(data.get("disclaimer")) or "以上内容用于健康管理沟通，不构成诊疗建议。如有具体健康问题，请线下就医。"
    document = f'''<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><title>{esc(customer)} - 体检解读与1v1方案</title>
<style>
@page {{ size: A4; margin: 18mm 16mm; }}
:root {{ --ink:#203238; --muted:#5f7074; --teal:#1d6b6a; --deep:#154b4c; --mint:#e9f3f0; --coral:#c65b48; --sand:#f7f5f0; --line:#d7e1df; }}
* {{ box-sizing:border-box; }} body {{ margin:0; color:var(--ink); background:#fff; font-family:-apple-system,BlinkMacSystemFont,"PingFang SC","Hiragino Sans GB","Microsoft YaHei",sans-serif; line-height:1.7; font-size:15px; }}
.page {{ max-width:840px; margin:0 auto; padding:34px 42px 56px; }}
.cover {{ min-height:980px; display:flex; flex-direction:column; justify-content:center; text-align:center; page-break-after:always; }}
.brand {{ color:var(--coral); font-weight:700; letter-spacing:.08em; }} h1 {{ color:var(--deep); font-size:42px; line-height:1.25; margin:20px 0 12px; }} h2 {{ color:var(--deep); font-size:28px; margin:42px 0 12px; border-bottom:2px solid var(--mint); padding-bottom:8px; }} h3 {{ color:var(--teal); font-size:19px; margin:28px 0 8px; }} p {{ margin:8px 0 12px; }} .sub {{ color:var(--muted); }}
.callout {{ border-left:5px solid var(--teal); background:var(--mint); padding:16px 20px; margin:16px 0 20px; }} .callout.sand {{ border-left-color:var(--coral); background:var(--sand); }} .callout strong {{ color:var(--teal); }} .callout.sand strong {{ color:var(--coral); }} .callout p {{ margin:5px 0 0; }}
.summary {{ margin:18px 0 24px; }} .summary-row {{ display:grid; grid-template-columns:96px 1fr; border-bottom:1px solid #fff; background:var(--sand); }} .summary-row span {{ background:var(--teal); color:#fff; padding:13px 15px; font-weight:700; }} .summary-row p {{ margin:0; padding:13px 17px; }}
.profile {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:8px; margin:12px 0 24px; }} .profile-item {{ background:var(--sand); padding:12px 15px; }} .profile-item span {{ display:block; color:var(--muted); font-size:12px; }} .profile-item strong {{ color:var(--deep); }}
ul {{ margin:8px 0 18px; padding-left:24px; }} li {{ margin:6px 0; }} table {{ width:100%; border-collapse:collapse; margin:12px 0 24px; font-size:14px; }} th {{ text-align:left; background:var(--deep); color:#fff; padding:11px 12px; }} td {{ vertical-align:top; border:1px solid var(--line); padding:11px 12px; }} tr:nth-child(even) td {{ background:var(--sand); }} td ul {{ margin:0; padding-left:18px; }} .page-break {{ page-break-before:always; }} .disclaimer {{ margin-top:28px; background:var(--sand); border-left:5px solid var(--coral); padding:14px 18px; color:var(--muted); }} .footer {{ margin-top:48px; text-align:center; color:var(--muted); font-size:12px; }}
@media print {{ .page {{ padding:0; }} .cover {{ min-height:250mm; }} .final-footer {{ display:none; }} }}
</style></head><body><main class="page">
<section class="cover"><div class="brand">{esc(brand)}</div><h1>体检报告解读<br>与 1v1 健康管理方案</h1><div class="sub">为 {esc(customer)} 定制 · {esc(report_date)}</div>{html_callout("这份报告想解决的事", text(data.get("cover_subtitle")) or "看懂身体现在的状态，并找到真正做得下去的下一步。", "sand")}<div class="footer">{esc(text(data.get("advisor")) or "健康管理顾问")}</div></section>
<section><h2>01 · 先看结论</h2><p class="sub">先把最重要的判断放在前面，避免被体检单上的一堆箭头带偏。</p>{profile_block}<div class="summary">{summary_html}</div></section>
<section><h2>02 · 体检报告解读</h2><p class="sub">{esc(report.get("intro"))}</p>{report_html}</section>
<section class="page-break"><h2>03 · 1v1 定制方案</h2><p class="sub">{esc(plan.get("intro"))}</p>{plan_html}</section>
<section><h2>04 · 说明</h2>{notes_html}<div class="disclaimer">{esc(disclaimer)}</div><div class="footer final-footer">{esc(brand)} · 客户专属健康管理资料 · {esc(report_date)}</div></section>
</main></body></html>'''
    output_path.write_text(document, encoding="utf-8")


def validate(data: dict[str, Any]) -> None:
    if not isinstance(data, dict):
        raise ValueError("输入 JSON 必须是对象")
    if not text(data.get("customer_name")):
        raise ValueError("缺少 customer_name")
    if not list_value(data.get("summary")):
        raise ValueError("缺少 summary，至少需要一条客户结论")
    if not isinstance(data.get("report", {}), dict):
        raise ValueError("report 必须是对象")
    if not isinstance(data.get("plan", {}), dict):
        raise ValueError("plan 必须是对象")


def main() -> int:
    parser = argparse.ArgumentParser(description="生成客户版体检解读与1v1健康管理方案")
    parser.add_argument("input", type=Path, help="内部结构化 JSON 文件")
    parser.add_argument("--out-dir", type=Path, default=None, help="输出目录，默认为输入文件所在目录")
    parser.add_argument("--date", default=None, help="文档日期，格式 YYYY-MM-DD")
    parser.add_argument("--pdf", action="store_true", help="同时生成 PDF；失败时保留 DOCX/HTML")
    args = parser.parse_args()

    data = json.loads(args.input.read_text(encoding="utf-8"))
    validate(data)
    report_date = text(args.date) or text(data.get("date")) or date.today().isoformat()
    out_dir = args.out_dir or args.input.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = f"{safe_name(text(data['customer_name']))}-{report_date}-体检解读与1v1方案"
    docx_path = out_dir / f"{stem}.docx"
    html_path = out_dir / f"{stem}.html"
    build_docx(data, docx_path, report_date)
    build_html(data, html_path, report_date)

    pdf_path = out_dir / f"{stem}.pdf"
    if args.pdf:
        try:
            with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                from weasyprint import HTML
        except (ImportError, OSError) as exc:
            print(
                "PDF 生成未成功，已保留 DOCX/HTML：PDF 引擎无法加载系统图形库。"
                "macOS 可运行 brew install pango；若仍失败，再检查 "
                "scripts/requirements-docs.txt。",
                file=sys.stderr,
            )
        else:
            try:
                HTML(filename=str(html_path), base_url=str(html_path.parent)).write_pdf(str(pdf_path))
            except Exception as exc:  # pragma: no cover - platform font/runtime failures
                print(f"PDF 生成未成功，已保留 DOCX/HTML：{exc}", file=sys.stderr)

    print(docx_path)
    print(html_path)
    if pdf_path.exists():
        print(pdf_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
